import smtplib
from datetime import datetime, timedelta

from django.core.files.storage import FileSystemStorage
from django.db.models import Sum
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect

# Create your views here.

from myapp.models import *



def login(request):
    return render(request,'loginindex.html')

def login_post(request):
    username=request.POST['textfield']
    password=request.POST['textfield2']
    l=Login.objects.filter(Username=username,Password=password)
    if l.exists():
        l2=Login.objects.get(Username=username,Password=password)
        if l2.Password==password:
            request.session['lid']=l2.id
            if l2.Type=='admin':
                # return HttpResponse('''<script>alert("Login Success");window.location="/myapp/adminHome/"</script>''')
                return render(request,"Admin/Adminindex.html")
            elif l2.Type=='staff':
                # return HttpResponse('''<script>alert("Login Success");window.location="/myapp/adminHome/"</script>''')
                return redirect('/myapp/staffHome/')
            elif l2.Type == 'wardmember':
                wm=Wardmember.objects.get(LOGIN__id=l2.id)
                ward=wm.WARD.id
                request.session['wid']=ward
                print(request.session['wid'])
                # return HttpResponse('''<script>alert("Login Success");window.location="/myapp/adminHome/"</script>''')
                return redirect('/myapp/memberHome/')
                # render(request,"staff/staffindex.html")
            else:
                return HttpResponse('''<script>alert("Login error");window.location="/myapp/login/"</script>''')

    return HttpResponse('''<script>alert("Login error");window.location="/myapp/login/"</script>''')




def adminHome(request):
    return render(request,'Admin/Adminindex.html')

def adminHome1(request):
    return render(request,'Admin/Adminindex1.html')

def staffHome(request):
    return render(request,'staff/staffindex.html')

def staffHome1(request):
    return render(request,'staff/staffindex1.html')

def memberHome(request):
    return render(request,'wardmember/memberindex.html')

def memberHome1(request):
    return render(request,'wardmember/memberindex1.html')


#======ADMIN=============


# department

def addDept(request):
    return render(request,'Admin/addDept.html')
def addDept_post(request):
    dname = request.POST['textfield']
    email = request.POST['textfield2']
    phone = request.POST['textfield22']
    if Department.objects.filter(Email=email).exists():
        return HttpResponse('''<script>alert("Department already exists");window.location="/myapp/viewDepartment/"</script>''')
    else:
      d=Department()
      d.Dname=dname
      d.Email=email
      d.Phone=phone
      d.save()
    return HttpResponse('''<script>alert("added");window.location="/myapp/viewDepartment/"</script>''')

def viewDepartment(request):
    obj=Department.objects.all()
    return  render(request,'Admin/manageDepartment.html',{'data':obj})

def viewDepartment_post(request):
    var=request.POST['textfield']
    obj=Department.objects.filter(Dname=var)
    return  render(request,'Admin/manageDepartment.html',{'data':obj})


def editDept(request,id):
    var= Department.objects.get(id=id)
    return render(request,'Admin/editDept.html',{'data':var})

def editDept_post(request):
    id=request.POST['id']
    dname = request.POST['textfield']
    email = request.POST['textfield2']
    phone = request.POST['textfield22']
    d = Department.objects.get(id=id)
    d.Dname = dname
    d.Email = email
    d.Phone = phone
    d.save()
    return HttpResponse('''<script>alert("Updated");window.location="/myapp/viewDepartment/"</script>''')

def deleteDept(request,id):
    var=Department.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewDepartment/"</script>''')

# staff

def addstaff(request):
    res=Department.objects.all()
    return render(request,'Admin/addstaff.html',{"data":res})

def addstaff_post(request):
    sname = request.POST['textfield']
    phone = request.POST['textfield1']
    email = request.POST['textfield2']
    gender = request.POST['radiobutton']
    DOB = request.POST['textfield3']
    Post = request.POST['textfield4']
    Designation = request.POST['textfield5']
    Pin = request.POST['textfield6']
    dept = request.POST['select']
    photo = request.FILES['file']

    fs=FileSystemStorage()
    date=datetime.now().strftime("%Y%m%d-%H%M%S")+'.jpg'
    fs.save(date,photo)
    path=fs.url(date)

    l=Login()
    l.Username=email
    l.Password=phone
    l.Type='staff'
    l.save()
    if Staff.objects.filter(Email=email).exists():
        return HttpResponse('''<script>alert("Staff already exists");window.location="/myapp/viewStaff/"</script>''')
    else:

      s=Staff()
      s.Sname=sname
      s.Phno=phone
      s.Email=email
      s.Gender=gender
      s.Dob=DOB
      s.Designation=Designation
      s.Post=Post
      s.Pin=Pin
      s.DEPARTMENT_id=dept
      s.Photo=path
      s.LOGIN=l
      s.save()
    return HttpResponse('''<script>alert("added");window.location="/myapp/viewStaff/"</script>''')



def viewStaff(request):
    obj=Staff.objects.all()
    obj2 = Department.objects.all()
    return  render(request,'Admin/manageStaff.html',{'data':obj,'data2':obj2,'sel':''})

def viewStaff_post(request):
    var=request.POST['textfield']
    var2=Department.objects.get(id=var)
    obj=Staff.objects.filter(DEPARTMENT_id=var)
    obj2=Department.objects.all()
    return  render(request,'Admin/manageStaff.html',{'data':obj,'data2':obj2,'sel':var,'seldept':var2.Dname})



def editStaff(request,id):
    var= Staff.objects.get(id=id)
    var2=Department.objects.all()
    return render(request,'Admin/editStaff.html',{'data':var,'data1':var2})

def editStaff_post(request):
    id=request.POST['id']
    sname = request.POST['textfield']
    phone = request.POST['textfield1']
    email = request.POST['textfield2']
    gender = request.POST['radiobutton']
    DOB = request.POST['textfield3']
    Post = request.POST['textfield4']
    Designation = request.POST['textfield5']
    Pin = request.POST['textfield6']
    dept = request.POST['select']
    obj = Staff.objects.get(id=id)

    if 'files' in request.FILES:
        photo = request.FILES['files']
        from datetime import datetime
        date = datetime.now().strftime('%Y%m%d-%H%M%S') + '.jpg'
        fs = FileSystemStorage()
        fn=fs.save(date, photo)
        obj.Photo = fs.url(date)
        obj.save()

    obj.Sname = sname
    obj.Phno = phone
    obj.Email = email
    obj.Gender = gender
    obj.Dob = DOB
    obj.Designation = Designation
    obj.Post = Post
    obj.Pin = Pin
    obj.DEPARTMENT_id = dept
    obj.save()
    return HttpResponse('''<script>alert("Updated");window.location="/myapp/viewStaff/"</script>''')


def deleteStaff(request,id):
    var=Staff.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewStaff/"</script>''')


# ward

def addWard(request):
    return render(request,'Admin/addward.html')


def addWard_post(request):
    wname = request.POST['textfield22']
    wnum = request.POST['select2']
    if Ward.objects.filter(Wnum=wnum).exists():
      return HttpResponse('''<script>alert("Ward already exists");window.location="/myapp/viewWard/"</script>''')
    else:
       w = Ward()
       w.Wname = wname
       w.Wnum = wnum
       w.save()
    return HttpResponse('''<script>alert("added");window.location="/myapp/viewWard/"</script>''')


def viewWard(request):
    obj=Ward.objects.all()
    return  render(request,'Admin/manageWard.html',{'data':obj})

def viewWard_post(request):
    var=request.POST['textfield']
    obj=Ward.objects.filter(Wnum=var)
    return  render(request,'Admin/manageWard.html',{'data':obj})

def deleteWard(request,id):
    var=Ward.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewWard/"</script>''')

def editWard(request,id):
    var= Ward.objects.get(id=id)
    return render(request,'Admin/editWard.html',{'data':var})

def editWard_post(request):
    id=request.POST['id']
    wname = request.POST['textfield22']
    wnum = request.POST['select2']
    w = Ward.objects.get(id=id)
    w.Wname = wname
    w.Wnum = wnum
    w.save()
    return HttpResponse('''<script>alert("Updated");window.location="/myapp/viewWard/"</script>''')



# wardmember

def addWardmember(request):
    obj=Ward.objects.all()
    return render(request,'Admin/addwardmember.html',{'data':obj})


def addWardmember_post(request):
    wname = request.POST['textfield1']
    phone = request.POST['textfield2']
    email = request.POST['textfield3']
    gender = request.POST['radiobutton']
    DOB = request.POST['textfield4']
    designation = request.POST['select2']
    Post = request.POST['textfield5']
    Pin = request.POST['textfield7']
    ward = request.POST['select']
    photo = request.FILES['file']

    fs = FileSystemStorage()
    date = datetime.now().strftime("%Y%m%d-%H%M%S") + '.jpg'
    fs.save(date, photo)
    path = fs.url(date)

    l = Login()
    l.Username = email
    l.Password = phone
    l.Type = 'wardmember'
    l.save()
    if Wardmember.objects.filter(Email=email).exists():
        return HttpResponse('''<script>alert("Wardmember already exists");window.location="/myapp/viewWardmember/"</script>''')
    else:

        s = Wardmember()
        s.Mname = wname
        s.Phno = phone
        s.Email = email
        s.Gender = gender
        s.Dob = DOB
        s.Designation = designation
        s.Post = Post
        s.Pin = Pin
        s.WARD_id =ward
        s.Photo = path
        s.LOGIN = l
        s.save()
        return HttpResponse('''<script>alert("updated");window.location="/myapp/viewWardmember/"</script>''')

def viewWardmember(request):
    obj=Wardmember.objects.all()
    obj2=Ward.objects.all()
    return  render(request,'Admin/mangeWardmember.html',{'data':obj, 'data2':obj2, 'sel':''})


def viewWardmember_post(request):
    var=request.POST['textfield']
    var2=Ward.objects.get(id=var)
    obj2=Ward.objects.all()
    obj=Wardmember.objects.filter(WARD_id=var)
    return  render(request,'Admin/mangeWardmember.html',{'data':obj, 'data2':obj2, 'sel':var, 'selward':var2.Wname})

def editWardmember(request,id):
    var= Wardmember.objects.get(id=id)
    var2=Ward.objects.all()
    return render(request,'Admin/editWardmember.html',{'data':var,'data1':var2})

def editWardmember_post(request):
    id=request.POST['id']
    wname = request.POST['textfield1']
    phone = request.POST['textfield2']
    email = request.POST['textfield3']
    gender = request.POST['radiobutton']
    DOB = request.POST['textfield4']
    Designation = request.POST['select2']
    Post = request.POST['textfield5']
    Pin = request.POST['textfield7']
    ward = request.POST['select']

    s = Wardmember.objects.get(id=id)

    if 'file' in request.FILES:
        photo = request.FILES['file']
        from datetime import datetime
        date = datetime.now().strftime('%Y%m%d-%H%M%S') + '.jpg'
        fs = FileSystemStorage()
        fn=fs.save(date, photo)
        s.Photo = fs.url(date)
        s.save()

    s.Mname = wname
    s.Phno = phone
    s.Email = email
    s.Gender = gender
    s.Dob = DOB
    s.Designation = Designation
    s.Post = Post
    s.Pin = Pin
    s.WARD_id = ward
    s.save()
    return HttpResponse('''<script>alert("Updated");window.location="/myapp/viewWardmember/"</script>''')

def deleteWardmember(request,id):
    var=Wardmember.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewWardmember/"</script>''')
#User

def viewUser(request):
    obj=User.objects.filter(Status="Approved")
    obj2=Ward.objects.all()
    return  render(request,'Admin/viewUser.html',{'data':obj, 'data2':obj2, 'sel':''})


# def viewUser_wardmem(request):
#     obj=Wardmember.objects.get(LOGIN_id=request.session['lid'])
#     return render(request, 'wardmember/memberindex.html', {'data': obj})

def viewUser_post(request):
    var=request.POST['textfield']
    var2=Ward.objects.get(id=var)
    obj2=Ward.objects.all()
    obj=User.objects.filter(WARD_id=var)
    return  render(request,'Admin/viewUser.html',{'data':obj, 'data2':obj2, 'sel':var, 'selward':var2.Wname})


# Certificate

def addCertificate(request):
    obj=Department.objects.all()
    return render(request,'Admin/addCertificate.html',{'data':obj})

def addCertificate_post(request):
    category = request.POST['category']
    desc = request.POST['description']
    dept = request.POST['select']
    if Certificate.objects.filter(Category=category).exists():
        return HttpResponse('''<script>alert("Service already exists");window.location="/myapp/viewCertificate/"</script>''')
    else:
      c=Certificate()
      c.Category=category
      c.Desc=desc
      c.DEPARTMENT_id=dept
      c.save()
    return HttpResponse('''<script>alert("certificate added");window.location="/myapp/viewCertificate/"</script>''')

def viewCertificate(request):
    var=Certificate.objects.all()
    return  render(request,'Admin/manageCertificates.html',{'data':var})

def viewCertificate_post(request):
    var=request.POST['textfield']
    obj=Certificate.objects.filter(Category=var)
    return  render(request,'Admin/manageCertificate.html',{'data':obj})

def editCertificate(request,id):
    var= Certificate.objects.get(id=id)
    obj = Department.objects.all()
    return render(request,'Admin/editCertificate.html',{'data':var,'data2':obj})

def editCertificate_post(request):
    id=request.POST['id']
    category = request.POST['category']
    desc = request.POST['description']
    dept = request.POST['select']
    cc = Department.objects.get(id=dept)
    c = Certificate.objects.get(id=id)
    c.Category = category
    c.Desc = desc
    c.DEPARTMENT_id = dept
    c.save()
    return HttpResponse('''<script>alert("certificate updated");window.location="/myapp/viewCertificate/"</script>''')


def deleteCertificate(request,id):
    var = Certificate.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewCertificate/"</script>''')

# scheme

def addScheme(request):
    obj=Department.objects.all()
    return render(request,'Admin/addScheme.html',{'data':obj})

def addScheme_post(request):
    scheme = request.POST['scheme']
    desc = request.POST['description']
    dept = request.POST['select']
    if Scheme.objects.filter(Scheme=scheme).exists():
        return HttpResponse('''<script>alert("Scheme already exists");window.location="/myapp/viewScheme/"</script>''')
    else:
       sc = Scheme()
       sc.Scheme = scheme
       sc.Desc = desc
       sc.DEPARTMENT_id = dept
       sc.save()
    return HttpResponse('''<script>alert("Scheme added");window.location="/myapp/viewScheme/"</script>''')

def viewScheme(request):
    obj=Scheme.objects.all()
    return  render(request,'Admin/manageSchemes.html',{'data':obj})

def viewScheme_post(request):
    var=request.POST['textfield']
    obj=Scheme.objects.filter(scheme=var)
    return  render(request,'Admin/manageSchemes.html',{'data':obj})

def editScheme(request,id):
    var= Scheme.objects.get(id=id)
    obj = Department.objects.all()
    return render(request,'Admin/editScheme.html',{'data':var,'data2':obj})

def editScheme_post(request):
    id=request.POST['id']
    scheme = request.POST['scheme']
    desc = request.POST['description']
    dept = request.POST['select']
    dd=Department.objects.get(id=dept)
    c = Scheme.objects.get(id=id)
    c.Scheme = scheme
    c.Desc= desc
    c.DEPARTMENT_id = dd
    c.save()
    return HttpResponse('''<script>alert(" updated");window.location="/myapp/viewScheme/"</script>''')


def deleteScheme(request,id):
    var = Scheme.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewScheme/"</script>''')


def admin_viewjobcardapplications(request,id):
    obj=Jobcardrequest.objects.filter(Status='verified', SCHEME_id=id)
    return render(request, 'Admin/jobcardrequest.html', {'data': obj})

def admin_viewjobcardapplications_post(request):
    var = request.POST['textfield']
    var2 = request.POST['textfield2']
    obj=Jobcardrequest.objects.filter(Status='verified', Date__range=[var,var2])
    return render(request, 'Admin/jobcardrequest.html', {'data': obj})


def admin_viewjobcarddetails(request,id):
    obj2=Jobcardmembers.objects.filter(JOBCARD__JOBCARDREQUEST_id=id)
    return render(request, 'Admin/jobcarddetails.html', {'data2':obj2})

def admin_approvejobcard(request,id):
    obj= Jobcardrequest.objects.get(id=id)
    Jobcardrequest.objects.filter(id=id).update(Status='approved')

    j = Jobcard.objects.get(JOBCARDREQUEST=obj)
    current_date = datetime.now().date()
    new_date = current_date + timedelta(days=5 * 365)
    j.Validity = new_date
    unique_cardno = generate_unique_cardno()
    j.Cardno = unique_cardno
    j.save()
    generate_pdf(request, id)
    return HttpResponse("<script>alert('approved');window.location='/myapp/admin_viewjobcardapplications/"+str(obj.SCHEME.id)+"'</script>")

def generate_unique_cardno():
    import uuid
    while True:
        # Generate a random UUID
        random_uuid = uuid.uuid4()
        # Format the UUID to match the desired pattern
        formatted_cardno = f"JC{random_uuid.hex[:2].upper()}-" \
                           f"{random_uuid.hex[8:12].upper()}-" \
                           f"{random_uuid.hex[12:16].upper()}-" \
                           f"{random_uuid.hex[:4].upper()}"
        # Check if the generated card number already exists in the database
        if not Jobcard.objects.filter(Cardno=formatted_cardno).exists():
            print(formatted_cardno)
            return formatted_cardno

def generate_pdf(request, id):
    from docx import Document
    from docx.shared import Inches
    from .models import Jobcardmembers
    import os
    from LGMS import settings
    from docx2pdf import convert


    obj2 = Jobcardmembers.objects.filter(JOBCARD__JOBCARDREQUEST_id=id)
    if len(obj2)>0:

        document = Document()

        document.add_heading('JOBCARD DETAILS', level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Card/No'
        hdr_cells[1].text = 'Validity'
        hdr_cells[2].text = 'Head'
        hdr_cells[3].text = 'Aadhar Number'
        hdr_cells[4].text = 'Photo'

        # i meant here
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = obj2[0].JOBCARD.Cardno
        hdr_cells[1].text = str(obj2[0].JOBCARD.Validity)
        hdr_cells[2].text = obj2[0].JOBCARD.JOBCARDREQUEST.Head
        hdr_cells[3].text = str(obj2[0].JOBCARD.JOBCARDREQUEST.Rationcard)

        try:
            photo_path = r'C:\Users\Hinil\PycharmProjects\LGMS' + str(obj2[0].JOBCARD.JOBCARDREQUEST.Photo).replace('/', '\\')
            # photo_path = os.path.join(settings.MEDIA_ROOT, str(member.Photo))
            hdr_cells[4].paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.0))
            # hdr_cells[2].text = 'Photo'

        except Exception as e:
            print(e, 'eeee')
            photo_path = r'C:\Users\Hinil\PycharmProjects\LGMS\myapp\static\Adminhome\assets\images\about.jpg'
            # photo_path = os.path.join(settings.MEDIA_ROOT, str(member.Photo))
            hdr_cells[4].paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.0))

        print()
        document.add_heading('MEMBER DETAILS', level=1)

        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Gender'
        hdr_cells[2].text = 'Age'
        hdr_cells[3].text = 'Adhar Number'
        hdr_cells[4].text = 'Relation'
        hdr_cells[5].text = 'Photo'

        # Add data to the table
        for member in obj2:
            row_cells = table.add_row().cells
            row_cells[0].text = member.Name
            row_cells[1].text = member.Gender
            row_cells[2].text = str(member.Age)
            row_cells[3].text = str(member.Adharno)
            row_cells[4].text = member.Relation

            if member.Photo:
                try:
                    photo_path = r'C:\Users\Hinil\PycharmProjects\LGMS'+str(member.Photo).replace('/','\\')
                    # photo_path = os.path.join(settings.MEDIA_ROOT, str(member.Photo))
                    row_cells[5].paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.0))
                except:
                    photo_path = r'C:\Users\Hinil\PycharmProjects\LGMS\myapp\static\Adminhome\assets\images\about.jpg'
                    # photo_path = os.path.join(settings.MEDIA_ROOT, str(member.Photo))
                    row_cells[5].paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.0))

        # Save the document
        document_path = os.path.join(settings.MEDIA_ROOT, f'cards/jobcard_{id}.docx')
        document.save(document_path)
        convert(f"media/cards/jobcard_{id}.docx")

        download_link = os.path.join(settings.MEDIA_URL, f'cards/jobcard_{id}.pdf')

    # return render(request, 'Admin/login.html', {'download_links': download_links})


def admin_rejectjobcard(request,id):
    obj= Jobcardrequest.objects.get(id=id)
    Jobcardrequest.objects.filter(id=id).update(Status='rejected')
    return HttpResponse("<script>alert('rejected');window.location='/myapp/admin_viewjobcardapplications/" + str(
        obj.SCHEME.id) + "'</script>")

def addPlans(request):
    return render(request,'Admin/addPlans.html')

def addPlans_post(request):
    plan = request.POST['plan']
    desc = request.POST['description']
    photo = request.FILES['file']
    date = datetime.now().strftime('%Y%m%d-%H%M%S') + '.jpg'
    fs = FileSystemStorage()
    fn = fs.save(date, photo)
    path = fs.url(date)
    if Plan.objects.filter(Plan=plan).exists():
        return HttpResponse('''<script>alert("Plan already exists");window.location="/myapp/viewPlans/"</script>''')
    else:
       p = Plan()
       p.Plan = plan
       p.Desc = desc
       p.Photo=path
       p.save()
    return HttpResponse('''<script>alert("Plan added");window.location="/myapp/viewPlans/"</script>''')

def viewPlans(request):
    obj=Plan.objects.all()
    return  render(request,'Admin/managePlans.html',{'data':obj})

def viewPlans_post(request):
    var=request.POST['textfield']
    obj=Plan.objects.filter(plan=var)
    return  render(request,'Admin/managePlans.html',{'data':obj})

def editPlans(request,id):
    var= Plan.objects.get(id=id)
    return render(request,'Admin/editPlans.html',{'data':var})

def editPlans_post(request):
    id=request.POST['id']
    plan = request.POST['plan']
    desc = request.POST['description']
    c = Plan.objects.get(id=id)


    if 'file' in request.FILES:
        photo = request.FILES['file']
        from datetime import datetime
        date = datetime.now().strftime('%Y%m%d-%H%M%S') + '.jpg'
        fs = FileSystemStorage()
        fn = fs.save(date, photo)
        c.Photo = fs.url(date)
        c.save()

    c.Plan = plan
    c.Desc= desc
    c.save()
    return HttpResponse('''<script>alert("updated");window.location="/myapp/viewPlans/"</script>''')

def deletePlans(request,id):
    var=Plan.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("deleted");window.location="/myapp/viewPlans/"</script>''')

def admin_viewPlanReq(request):
    var=Planrequest.objects.filter(Status='pending')
    return render(request, 'Admin/viewPlanReq.html', {'data': var})

def admin_approvePlanReq(request,id):
    obj =Planrequest.objects.filter(id=id).update(Status='approved')
    return HttpResponse('''<script>alert("Approved");window.location="/myapp/viewPlans/"</script>''')

def admin_viewApprovedPlans(request):
    var=Planrequest.objects.filter(Status='approved')
    return render(request, 'Admin/viewApprovedplans.html', {'data': var})



def admin_viewBirthapplications(request,id):
    obj=Birthcertificaterequest.objects.filter(Status='verified', BIRTH__CERTIFICATE_id=id)
    return render(request, 'Admin/viewBirthapplications.html', {'data': obj})

def admin_viewBirthapplications_post(request):
    var = request.POST['textfield']
    var2 = request.POST['textfield2']
    obj=Birthcertificaterequest.objects.filter(Status='verified', Date__range=[var,var2])
    return render(request, 'Admin/viewBirthapplications.html', {'data': obj})


def admin_approveBirthCertificate(request,id):
    obj= Birthregistration.objects.get(id=id)
    Birthregistration.objects.filter(id=id).update(Status='approved')
    return render(request, 'Admin/viewBirthapplications.html', {'data': obj})


def admin_viewBirthdetails(request,id):
    obj2=Birthregistration.objects.filter(id=Birthcertificaterequest.objects.get(id=id).BIRTH.id)
    return render(request, 'Admin/viewBirthDetails.html', {'data2':obj2})


# cirtificate generation

def admin_approveCertificate(request,id):
    obj= Birthregistration.objects.get(id=id)
    Birthregistration.objects.filter(id=id).update(Status='approved')
    Birthcertificaterequest.objects.filter(BIRTH_id=id).update(Status='approved')
# def generate_birthcertificate(request,id):


    from docx import Document
    from docx.shared import Pt
    from docx2pdf import convert

    def generate_certificate_docx():
        document = Document()

        # Define styles
        company_heading_style = document.styles['Heading 1']
        company_heading_style.font.bold = True
        company_heading_style.font.size = Pt(18)


        company_subheading_style = document.styles['Normal']
        company_subheading_style.font.size = Pt(12)

        content_style = document.styles['Normal']
        content_style.font.size = Pt(12)

        sign_style = document.styles['Heading 4']
        sign_style.font.bold = True
        sign_style.font.size = Pt(12)

        # Add content to the document
        # document.add_heading('GOVERNMENT OF KERALA', level=1)
        # document.add_heading('DEPARTMENT OF URBAN AFFAIRS', level=1)
        document.add_heading('BIRTH CERTIFICATE', level=0)
        document.add_paragraph(f'Register Number : {regno}', style='Heading 4')
        document.add_paragraph(f'This is to certify that {name} ({gender}) was born to  {father} and  {mother} on {dob} at {hospital}  in {district},kerala.', style='Normal')
        # document.add_paragraph(f"", style='Normal')
        # document.add_paragraph(f" {father} and", style='Normal')
        # document.add_paragraph(f" {mother}", style='Normal')
        # document.add_paragraph(f"on {dob}", style='Normal')
        # document.add_paragraph(f'at {hospital}', style='Normal')
        # document.add_paragraph(f' in {district,"kerala"}.', style='Normal')
        document.add_paragraph(f'Signature of Medical Officer', style='Heading 4')
        photo_path = r'C:\Users\Hinil\PycharmProjects\LGMS\myapp\static\Adminhome\assets\images\signature.jpg'
        from docx.shared import Inches
        document.add_picture(photo_path, width=Inches(1.0))
        # Save the document
        document.save(f'media/certificates/birth_certificate{obj.id}.docx')
        convert(f"media/certificates/birth_certificate{obj.id}.docx")



    name = obj.Name
    gender =obj.Gender
    dob =obj.Dob
    mother =obj.Mother
    father =obj.Father
    hospital =obj.Hospital
    regno =obj.Regno
    district =obj.Dist

    generate_certificate_docx()

    return render(request, 'Admin/viewBirthapplications.html')





#
# from reportlab.lib.pagesizes import letter
# from reportlab.lib import colors
# from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer, Paragraph
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from io import BytesIO
#
#
# def generate_certificate_pdf(date, user_name, user_phone, products, total_amount, bill_id):
#     buffer = BytesIO()
#     document = SimpleDocTemplate(buffer, pagesize=letter)
#     styles = getSampleStyleSheet()
#
#     company_heading_style = ParagraphStyle(
#         name='CompanyHeading',
#         parent=styles['Heading1'],
#         alignment=1,
#         fontName='Helvetica-Bold',
#         fontSize=18,
#     )
#     company_subheading_style = ParagraphStyle(
#         name='CompanySubHeading',
#         parent=styles['Normal'],
#         alignment=1,
#         fontName='Helvetica',
#         fontSize=12,
#         spaceAfter=10
#     )
#     content_style = ParagraphStyle(
#         name='Content',
#         parent=styles['Normal'],
#         alignment=0,
#         fontName='Helvetica',
#         fontSize=12,
#         leading=15,
#
#     )
#     sign_style = ParagraphStyle(
#         name='Content',
#         # parent=styles['Normal'],
#         # alignment=0,
#         # fontName='Helvetica',
#         fontSize=12,
#         leading=15,
#
#         parent=styles['Heading4'],  # Inherit from an existing style if needed
#         alignment=2,  # Center alignment
#         fontName='Helvetica-Bold',  # Specify font name and bold
#
#     )
#     cer_heading1 = Paragraph("<b>GOVERNMENT OF KERALA</b>", company_heading_style)
#     cer_heading2 = Paragraph("<b>DEPARTMENT OF URBAN AFFAIRS</b>", company_heading_style)
#     cer_heading3 = Paragraph("<b> BIRTH CERTIFICATE </b>", company_heading_style)
#
#     cer_subheading = Paragraph("<b>This Certificate shows that</b>", company_subheading_style)
#     content = Paragraph(
#         "<b><u>{Name}</u></b> Sex <b><u>Female</u> </b> was born at <b><u>{hospital}</u></b> in Keralaon the <b><u>{date}</u></b> day of<b> <u>{month ,year}</u></b> and that the parents names are as follows:",
#         content_style)
#     father_name = Paragraph("Father's name <b><u> {father name}</u></b>", content_style)
#     mother_name = Paragraph("Mother's name <b><u> {mother name}</u></b>", content_style)
#
#     date = Paragraph("Date  : {Date}", sign_style)
#     signature = Paragraph("{signature}", sign_style)
#     signature1 = Paragraph(f"Signature of Issuing Authority", sign_style)
#
#     date = Paragraph("Date  : <b>{Date}</b>", styles['Heading4'])
#
#     #
#     flowables = [cer_heading1, Spacer(1, 12), cer_heading2, Spacer(1, 12), cer_heading3, Spacer(5, 20), cer_subheading,
#                  content, Spacer(1, 12), father_name, Spacer(1, 12), mother_name, Spacer(3, 12), date, Spacer(1, 12),
#                  signature, signature1]
#     # flowables = [cer_heading1,cer_heading2,cer_heading3, bill_subheading,bill_subheading1,bill_subheading2,bill_subheading3, spacer,bill_id,signature ,spacer,customer_info, spacer, table, spacer, total_text,
#     #              spacer,spacer, thank_you_message]
#     document.build(flowables)
#
#     buffer.seek(0)
#     return buffer


#=========USER==========

def user_signup(request):
    name = request.POST['name']
    phone = request.POST['phone']
    email = request.POST['email']
    gender = request.POST['gender']
    DOB = request.POST['dob']
    house = request.POST['house']
    place = request.POST['place']
    Post = request.POST['post']
    Pin = request.POST['pin']
    sc = request.POST['sc']
    status = "pending"
    password= request.POST['password']
    confirmpassword= request.POST['confirmpassword']
    photo = request.POST['photo']
    ward = request.POST['ward']
    print(ward)
    print(password)
    print(confirmpassword)


    if Login.objects.filter(Username=email).exists():
        return JsonResponse({'status':'no'})
    l = Login()
    l.Username = email
    l.Password = password
    l.Type = 'user'
    l.save()

    if password == confirmpassword:

        date = datetime.now().strftime("%Y%m%d-%H%M%S") + '.jpg'
        import base64
        b = base64.b64decode(photo)
        open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\\'+date, 'wb').write(b)
        path = '/media/'+date

        s = User()
        s.Uname = name
        s.Phno = phone
        s.Email = email
        s.Gender = gender
        s.Dob = DOB
        s.Hname = house
        s.Post = Post
        s.Pin = Pin
        s.Place = place
        s.Sc = sc
        s.WARD_id = ward
        s.Photo = path
        s.Status=status
        s.LOGIN = l
        s.save()
        return JsonResponse({'status': 'ok'})

def user_login(request):
    username=request.POST['username']
    password=request.POST['password']
    log=Login.objects.filter(Username=username,Password=password)
    if log.exists():
        log1=Login.objects.get(Username=username,Password=password)
        lid=log1.id
        if log1.Type=='user':
            return JsonResponse({'status':'ok','lid':str(lid),'Type':log1.Type})
        elif log1.Type=='Employeemate':
            return JsonResponse({'status': 'ok', 'lid': str(lid), 'Type': log1.Type})
        else:
            return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status': 'no'})


def GetWard(request):
    w=Ward.objects.all()
    wa=[]
    for i in w:
        wa.append({'id':i.id,'Ward':i.Wname})
    return JsonResponse({'status':'ok','data':wa})


def viewProfile(request):
    lid=request.POST['lid']
    var=User.objects.get(LOGIN_id=lid)
    return JsonResponse({'status':'ok','Uname':var.Uname,
                         'Phno':var.Phno,
                         'Email':var.Email,
                         'Gender':var.Gender,
                         'Dob': var.Dob,
                         'Photo': var.Photo,
                         'Post': var.Post,
                         'Place':var.Place,
                         'Hname':var.Hname,
                         'Pin':var.Pin,
                         'WARD':var.WARD.Wname,
                         'Sc':var.Sc})





def editProfile(request):
    name = request.POST['name']
    ph_no = request.POST['phone']
    email = request.POST['email']
    dob = request.POST['dob']
    place = request.POST['place']
    post = request.POST['post']
    pin = request.POST['pin']
    hname = request.POST['house']
    sc = request.POST['sc']
    gender = request.POST['gender']
    photo = request.POST['photo']
    ward = request.POST['ward']

    lid = request.POST['lid']

    obj = User.objects.get(LOGIN_id=lid)
    if len(photo)>0:
        date = datetime.now().strftime("%Y%m%d-%H%M%S") + '.jpg'
        import base64
        b = base64.b64decode(photo)
        open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\\' + date, 'wb').write(b)
        path = '/media/' + date
        obj.Photo=path

    obj.Uname = name
    obj.Phno = ph_no
    obj.Email = email
    obj.Dob = dob
    obj.Gender = gender
    obj.Hname=hname
    obj.Place = place
    obj.Post = post
    obj.Pin = pin
    obj.Sc = sc
    obj.WARD_id=ward
    obj.save()


    lobj = Login.objects.get(id=lid)
    lobj.username = email
    lobj.save()



    return JsonResponse({'status':"ok"})





def UserViewServices(request):
    c=Certificate.objects.all()
    el=[]
    for i in c:
        el.append({'id':i.id,'Category':i.Category,
                   'Desc':i.Desc,
                   'DEPARTMENT':i.DEPARTMENT.Dname,
                   })
    return JsonResponse({'status':'ok','data':el})


def UserViewSchemes(request):
    c=Scheme.objects.all()
    el=[]
    for i in c:
        el.append({'id':i.id,'Scheme':i.Scheme,
                   'Desc':i.Desc,
                   'DEPARTMENT':i.DEPARTMENT.Dname,
                   })
    return JsonResponse({'status':'ok','data':el})




def jobcard(request):
    headname = request.POST['headname']
    rationcard = request.POST['rationcard']
    idproof=request.POST['idp']
    Photo = request.POST['photo']
    lid=request.POST['lid']
    sid=request.POST['sid']
    date = datetime.now().strftime("%Y%m%d-%H%M%S") + '.jpg'
    import base64
    b = base64.b64decode(Photo)
    open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\\' + date, 'wb').write(b)
    path = '/media/' + date

    date1 = datetime.now().strftime("%Y%m%d-%H%M%S") + '-1.jpg'
    import base64
    b1 = base64.b64decode(idproof)
    open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\\' + date1, 'wb').write(b1)
    path1 = '/media/' + date1

    res=Jobcardrequest.objects.filter(USER__LOGIN_id=lid)
    if res.exists():
        return JsonResponse({'status':'no'})
    else:
        b=Jobcardrequest()
        b.Head=headname
        b.Rationcard=rationcard
        b.Idproof=path1
        b.Photo=path
        b.Date=datetime.now().today()
        b.Status='pending'
        b.SCHEME_id=sid
        b.USER=User.objects.get(LOGIN_id=lid)
        b.save()


        jb=Jobcard()
        jb.Head=headname
        jb.Cardno=0
        jb.Rationcard=rationcard
        jb.Photo=path
        jb.Validity=datetime.now().date()
        jb.Date=datetime.now().today()
        jb.JOBCARDREQUEST=b
        jb.USER=User.objects.get(LOGIN_id=lid)
        jb.save()

        return JsonResponse({'status':'ok'})



def addmembers(request):
    name = request.POST['name']
    gender = request.POST['gender']
    Photo = request.POST['photo']
    jobcardId = request.POST['jobcardId']
    age = request.POST['age']
    relation = request.POST['relation']
    adharno = request.POST['adharno']
    bankname = request.POST['bankname']
    lid=request.POST['lid']
    sid=request.POST['sid']
    print(sid, lid, jobcardId)


    import base64
    b = base64.b64decode(Photo)
    date=datetime.now().strftime('%Y%m%d%H%M%S%f')+'.jpg'
    open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\\' + str(date), 'wb').write(b)
    path = '/media/' + str(date)

    b=Jobcardmembers()
    b.Name=name
    b.Gender=gender
    b.Age=age
    b.Relation=relation
    b.Adharno=adharno
    b.Bank=bankname
    b.Photo=path
    b.JOBCARD_id=Jobcard.objects.get(JOBCARDREQUEST_id=jobcardId).id
    b.save()
    return JsonResponse({'status':'ok'})



def viewJobcardDetails(request):
    lid = request.POST['lid']
    var = Jobcardrequest.objects.filter(USER__LOGIN_id=lid)
    card = ''
    if len(var)>0:
        var = var[0]
        a="no"
        if Workrequest.objects.filter(JOBCARD__USER__LOGIN_id=lid).exists():
            a="yes"
        print(var)
        if Jobcard.objects.filter(JOBCARDREQUEST=var).exists():
            jc = Jobcard.objects.filter(JOBCARDREQUEST=var)[0]
            if jc.Cardno!=0:
                card = '/media/cards/jobcard_'+str(var.id)+'.pdf'
        print(a,"mmmmmmmm,m")
        return JsonResponse({'status': 'ok', 'Head': var.Head,'Idproof':var.Idproof,
                             'Rationcard': var.Rationcard,
                             'photo': var.Photo, 'jobcardId':var.id, 'card':card, 'Status':var.Status,'a':a})

    return JsonResponse({'status': 'ok', 'Head': 'None','Idproof':'',
                                 'card': '',
                                 'Rationcard': '',
                                 'photo': '', 'jobcardId':'', 'Status':'','a':''})



def Viewmembers(request):
    jobcardId=request.POST['jobcardId']
    lid=request.POST['lid']
    c=Jobcardmembers.objects.filter(JOBCARD__JOBCARDREQUEST_id=jobcardId)
    # if EmployeemateReq.objects.filter()
    el=[]
    for i in c:
        el.append({'id':i.id,'Name':i.Name,
                   'Gender':i.Gender,
                   'Age':i.Age,
                   'Adharno':i.Adharno,
                   'Photo':i.Photo,
                   'Relation':i.Relation,
                   })
    print(el)
    return JsonResponse({'status':'ok','data':el})




def UserViewNotifications(request):
    lid = request.POST['lid']
    c=PlanNotification.objects.filter(WARD__user__LOGIN_id=lid).order_by('-Date')
    el=[]
    for i in c:
        el.append({'id':i.id,'Title':i.title,
                   'Desc':i.Desc,
                   'Photo':i.PLANREQUEST.PLAN.Photo,
                   'Sdate':i.PLANREQUEST.Sdate,
                   'Edate':i.PLANREQUEST.Edate,
                   })
    return JsonResponse({'status':'ok','data':el})







def BirthRegister(request):
    lid = request.POST['lid']
    cid = request.POST['cid']
    name = request.POST['name']
    gender = request.POST['gender']
    dob = request.POST['dob']
    district = request.POST['district']
    mother = request.POST['mother']
    father = request.POST['father']
    hospital = request.POST['hospital']
    regno = request.POST['regno']

    c=Birthregistration()
    c.Dist=district
    c.Name=name
    c.Gender=gender
    c.Dob=dob
    c.Mother=mother
    c.Father=father
    c.Hospital = hospital
    c.Regno=regno
    c.Status = 'pending'
    c.CERTIFICATE_id = cid
    c.USER = User.objects.get(LOGIN_id=lid)
    c.save()
    print(lid,"lid")
    print(cid,'certificate')

    br=Birthcertificaterequest()
    br.USER=User.objects.get(LOGIN_id=lid)
    br.Date=datetime.now().today()
    br.BIRTH=c
    br.Status='pending'
    br.Bcertificate='pending'
    br.save()

    return JsonResponse({'status':'ok'})



# def viewDeathapplications(request):
#     # obj=Jobcardrequest.objects.filter(Status='pending',jobcard__in=Jobcardmembers.objects.filter().values_list('JOBCARD'))
#     obj = Deathregistration.objects.filter(Status='pending')
#     return render(request, 'Staff/Deathapplications.html', {'data': obj})


def viewBirthapplications(request,id):
    obj=Birthcertificaterequest.objects.filter(Status='pending', BIRTH__CERTIFICATE_id=id)
    return render(request, 'Staff/Birthapplications.html', {'data': obj})
#
# def viewBirthapplications(request):
#     obj=Birthcertificaterequest.objects.filter(Status='pending')
#     return render(request, 'Staff/Birthapplications.html', {'data': obj})

def viewBirthapplications_post(request):
    var = request.POST['textfield']
    var2 = request.POST['textfield2']
    obj=Birthcertificaterequest.objects.filter()
    return render(request, 'Staff/Birthapplications.html', {'data': obj})


def viewBirthdetails(request,id):
    obj2=Birthregistration.objects.filter(id=Birthcertificaterequest.objects.get(id=id).BIRTH.id)
    return render(request, 'Staff/BirthDetails.html', {'data2':obj2})

#
# def viewBirthdetails_post(request):
#     var = request.POST['textfield']
#     obj2=Birthregistration.objects.filter()
#     return render(request, 'Staff/BirthDetails.html', {'data2':obj2})

def staff_verifyBirthapplications(request,id):
    obj= Birthcertificaterequest.objects.get(BIRTH_id=id)
    Birthcertificaterequest.objects.filter(BIRTH_id=id).update(Status='verified')
    Birthregistration.objects.filter(id=obj.BIRTH.id).update(Status='verified')
    return HttpResponse("<script>alert('Verified');window.location='/myapp/viewBirthapplications/"+str(id)+"'</script>")


def staff_rejectBirthapplications(request,id):
    obj= Birthcertificaterequest.objects.get(BIRTH_id=id)
    var=Birthcertificaterequest.objects.filter(BIRTH_id=id).update(Status='Rejected')
    Birthregistration.objects.filter(id=obj.BIRTH.id).update(Status='Rejected')
    return HttpResponse('''<script>alert("Rejected");window.location="/myapp/viewBirthapplications/"</script>''')


def viewmycertificates(request):
    lid =  request.POST['lid']
    obj=Birthcertificaterequest.objects.filter(USER__LOGIN_id=lid)
    el = []
    for i in obj:
        el.append({'id': i.BIRTH.id, 'Date': i.Date,
                   'regno': i.BIRTH.Regno,
                   'status': i.Status,
                   })
    return JsonResponse({'status': 'ok', 'data': el})


def DeathRegister(request):
    lid = request.POST['lid']
    cid = request.POST['cid']
    name = request.POST['name']
    gender = request.POST['gender']
    district = request.POST['district']
    dod = request.POST['dod']
    dob = request.POST['dob']
    wife = request.POST['wife']
    husband = request.POST['husband']
    hospital = request.POST['hospital']
    regno = request.POST['regno']

    u=Deathregistration()
    u.Name = name
    u.Gender = gender
    u.Dist = district
    u.Dod = dod
    u.Dob = dob
    u.Mother = wife
    u.Father = husband
    u.Hospital = hospital
    u.Regno = regno
    u.Status = 'pending'
    u.CERTIFICATE_id = cid
    u.USER = User.objects.get(LOGIN_id=lid)
    u.save()
    return JsonResponse({'status': 'ok'})

#======STAFF=============

def staffProfile(request):
    res=Staff.objects.get(LOGIN_id=request.session['lid'])
    return render(request,'staff/viewProfile.html',{'data':res})


def staffviewScheme(request):
    obj=Scheme.objects.all()
    return  render(request,'Staff/Schemes.html',{'data':obj})

def staffviewScheme_post(request):
    var=request.POST['textfield']
    obj=Scheme.objects.filter(scheme=var)
    return  render(request,'Staff/Schemes.html',{'data':obj})


def staffviewCertificate(request):
    var=Certificate.objects.all()
    return  render(request,'Staff/Services.html',{'data':var})

def staffviewCertificate_post(request):
    var=request.POST['textfield']
    obj=Certificate.objects.filter(Category=var)
    return  render(request,'Staff/Services.html',{'data':obj})


def viewjobcards(request,id):
    obj=Jobcardrequest.objects.filter(Status='verified',jobcard__in=Jobcardmembers.objects.filter().values_list('JOBCARD'))
    obj2=Jobcard.objects.filter(JOBCARDREQUEST__SCHEME_id=id)
    return render(request, 'Staff/jobcards.html', {'data': obj,'data2': obj2})

#
# def viewjobcards_post(request,id):
#     obj=Jobcardrequest.objects.filter(Status='verified',jobcard__in=Jobcardmembers.objects.filter().values_list('JOBCARD'))
#     obj2=Jobcard.objects.all()
#     return render(request, 'Staff/jobcards.html', {'data': obj,'data2': obj2})
#


def viewjobcardapplications(request):
    # obj=Jobcardrequest.objects.filter(Status='pending',jobcard__in=Jobcardmembers.objects.filter().values_list('JOBCARD'))
    obj=Jobcardrequest.objects.filter(Status='pending')
    return render(request, 'Staff/jobcardApplications.html', {'data': obj})

def viewjobcardapplications_post(request):
    var = request.POST['textfield']
    var2 = request.POST['textfield2']
    obj=Jobcardrequest.objects.filter()
    return render(request, 'Staff/jobcardApplications.html', {'data': obj})


def viewjobcarddetails(request,id):
    obj2=Jobcardmembers.objects.filter(JOBCARD__JOBCARDREQUEST_id=id)
    return render(request, 'Staff/jobcarddetails.html', {'data2':obj2})


def viewjobcarddetails_post(request):
    var = request.POST['textfield']
    obj2=Jobcardmembers.objects.filter()
    return render(request, 'Staff/jobcarddetails.html', {'data2':obj2})

def staff_verifyjobcard(request,id):
    obj= Jobcardrequest.objects.get(id=id)
    Jobcardrequest.objects.filter(id=id).update(Status='verified')
    return HttpResponse("<script>alert('Verified');window.location='/myapp/viewjobcards/"+str(obj.SCHEME.id)+"'</script>")


def rejectjobcard(request,id):
    var=Jobcardrequest.objects.filter(id=id).update(Status='Rejected')
    return HttpResponse('''<script>alert("Rejected");window.location="/myapp/viewjobcardapplications/"</script>''')
#======WARD MEMBER=============


def memberProfile(request):
    res=Wardmember.objects.get(LOGIN_id=request.session['lid'])
    return render(request,'wardmember/viewProfile.html',{'data':res})

#
def verifyUser(request):
    obj=User.objects.filter(WARD__id=request.session['wid'],Status="Pending")
    print(obj,"hhhhhhhhhh")
    return  render(request,'wardmember/verifyUser.html',{'data':obj, 'sel':''})

def verifyUser_post(request):
    var=request.POST['textfield']
    obj=User.objects.filter(Uname=var)
    return  render(request,'wardmember/verifyUser.html',{'data':obj})


def approve(request,id):
    var=User.objects.filter(LOGIN=id).update(Status='Approved')
    v=User.objects.filter(LOGIN=id)[0]
    return HttpResponse('''<script>alert("Approved");window.location="/myapp/verifyUser/"</script>''')

def reject(request,id):
    var=User.objects.filter(LOGIN=id).update(Status='Rejected')
    return HttpResponse('''<script>alert("Rejected");window.location="/myapp/verifyUser/"</script>''')

def acceptedUser(request):
    var=User.objects.filter(Status='Approved')
    return render(request,"wardmember/acceptedusers.html",{'data':var})

def acceptedUser_post(request):
    var=request.POST['textfield']
    obj=User.objects.filter(Uname=var)
    return  render(request,'wardmember/acceptedusers.html',{'data':obj})


def viewWardPlans(request):
    obj=Plan.objects.all()
    l = []
    for i in obj:
        status = ''
        noti = ''
        rid=''
        if Planrequest.objects.filter(PLAN=i,MEMBER__LOGIN_id=request.session['lid']).exists():
            vars=Planrequest.objects.filter(PLAN=i,MEMBER__LOGIN_id=request.session['lid'])[0]
            status=vars.Status
            rid=vars.id
            var2=PlanNotification.objects.filter(PLANREQUEST=vars)
            if var2.exists():
                noti='Active'
        l.append({
            'id':i.id,
            'Plan':i.Plan,
            'Photo':i.Photo,
            'Desc':i.Desc,
            'status':status,
            'rid':rid,
            'noti':noti,
        })
    return  render(request,'Wardmember/viewPlans.html',{'data':l})

def viewWardPlans_post(request):
    var=request.POST['textfield']
    obj=Plan.objects.filter(plan=var)
    return  render(request,'Wardmember/viewPlans.html',{'data':obj})

def sendPlanReq(request,id):
    var=str(datetime.now().date())
    return  render(request,'Wardmember/planRequest.html',{'id':id,'dt':var})

def sendPlanReq_post(request):
    start = request.POST['sdate']
    end = request.POST['edate']
    id=request.POST['id']
    if Planrequest.objects.filter(PLAN_id=id,MEMBER__LOGIN_id=request.session['lid']).exists():
        return HttpResponse("<script>alert('Request Already Exists');history.back()</script>")

    p = Planrequest()
    p.Sdate = start
    p.Edate = end
    p.Status='pending'
    p.PLAN_id=id
    ww=Wardmember.objects.get(LOGIN=request.session['lid'])
    p.MEMBER_id=ww.id
    p.save()
    return HttpResponse('''<script>alert("Requested");window.location="/myapp/viewWardPlans/"</script>''')

def sendPlanNotification(request,id):
    return  render(request,'Wardmember/sendPlanNoti.html',{'id':id})


def sendPlanNotification_post(request):
    id=request.POST['id']
    titlee = request.POST['title']
    desc = request.POST['description']

    p = PlanNotification()
    p.title = titlee
    p.Desc = desc
    p.PLANREQUEST_id=id
    p.Date=datetime.now().today().date()
    p.WARD=Wardmember.objects.get(LOGIN=request.session['lid']).WARD
    p.save()
    return HttpResponse('''<script>alert("Sended");window.location="/myapp/viewWardPlans/"</script>''')


def logout(request):
    request.session['lid']=''
    return redirect('/myapp/login/')




#========EMPLOYEEMATE===========



def sendEmployeematereq(request):
    sslc=request.POST['sslc']
    email=request.POST['email']
    lid=request.POST['lid']
    mid=request.POST['mid']
    jobcardmemId=request.POST['jobcardmemId']

    import base64
    b = base64.b64decode(sslc)
    date = datetime.now().strftime('%Y%m%d%H%M%S%f') + '.jpg'
    open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\images\\' + str(date), 'wb').write(b)
    path = '/media/images/' + str(date)

    print(jobcardmemId,'jjjjjjj')

    # import random
    #
    # objj=Login()
    # objj.Username=email
    # objj.Password=random.randint(0000,9999)
    # objj.Type='Employeemate'
    # objj.save()
    if EmployeemateReq.objects.filter(USER=User.objects.get(LOGIN_id=lid),JOBCARDMEMBERS_id=mid).exists():
        return JsonResponse({'status': 'no'})

    obj=EmployeemateReq()
    obj.Sslc=path
    obj.USER=User.objects.get(LOGIN_id=lid)
    obj.Status='pending'
    obj.JOBCARDMEMBERS_id=mid
    obj.LOGIN_id=1
    obj.Email=email
    obj.save()
    return JsonResponse({'status':'ok'})



def viewEmployeemateApplication(request):
    obj = EmployeemateReq.objects.filter(Status='pending')
    return render(request, 'Staff/EmployeemateRequest.html', {'data': obj})


def viewEmployeemates(request):
    obj2= EmployeemateReq.objects.filter(Status='verified')
    return render(request, 'Staff/Employeemates.html', {'data2':obj2})

def viewEmployeemates_post(request):
    var=request.POST['textfield']
    obj = EmployeemateReq.objects.filter(Status='verified')
    return  render(request,'Staff/EmployeemateRequest.html',{'data':obj})

def verifyEmployeemateApplication(request,id):
    obj = EmployeemateReq.objects.filter(id=id).update(Status='verified')
    objj=EmployeemateReq.objects.get(id=id).Email

    l=Login()
    l.Username=objj

    import random
    new_pass = random.randint(000000, 999999)
    l.Password=str(new_pass)
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login("projectlgms916@gmail.com", "rykk yaow wbwi twok")  # App Password
    to = objj
    subject = "Test Email"
    body = "Your  password is " + str(new_pass)+"\n"+ "Your Username is " + str(objj)
    msg = f"Subject: {subject}\n\n{body}"
    server.sendmail("s@gmail.com", to, msg)
    # Disconnect from the server
    server.quit()
    # ress = Login.objects.filter(username=objj).filter(password=new_pass)
    l.Type='Employeemate'
    l.save()

    emt=Employeemate()
    emt.Date=datetime.now().today()
    emt.Status='Verified'
    emt.JOBCARD_id=id
    emt.LOGIN_id=l.id
    emt.EMPLOYEEMATEREQ_id=id
    current_date = datetime.now().date()
    one_year_from_now = current_date + timedelta(days=365)

    emt.validity=one_year_from_now
    emt.save()


    return HttpResponse(
            '''<script>alert('Verified');window.location='/myapp/viewEmployeemateApplication/'</script>''')

    # Disconnect from the serverserver.quit()

def rejectEmpmateReq(request,id):
    var=EmployeemateReq.objects.filter(id=id).update(Status='Rejected')
    return HttpResponse('''<script>alert("Rejected");window.location="/myapp/viewEmployeemateApplication/"</script>''')


def viewEmployeemateProfile(request):
    lid=request.POST['lid']
    var=Employeemate.objects.get(LOGIN_id=lid)
    return JsonResponse({'status':'ok','Name':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.Name,
                         'Gender':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.Gender,
                         'Photo':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.Photo,
                         'Age':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.Age,
                         'Cardno':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.Cardno,
                         'ward':var.EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD.Wname,
                         'Email':var.EMPLOYEEMATEREQ.Email,})


def deleteEmpmate(request,id):
    var=EmployeemateReq.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("Deleted");window.location="/myapp/viewEmployeemates/"</script>''')


def useraddapplyworkreq(request):
    lid=request.POST['lid']
    # sid=request.POST['sid']
    sid=1
    area=request.POST['area']
    taxrecipt=request.POST['photo']
    from datetime import datetime
    current_date=datetime.now().today()
    import base64
    b = base64.b64decode(taxrecipt)
    date = datetime.now().strftime('%Y%m%d%H%M%S%f') + '.jpg'
    open(r'C:\Users\Hinil\PycharmProjects\LGMS\media\taxrecpt\\' + str(date), 'wb').write(b)
    path = '/media/taxrecpt/' + str(date)

    # new_date = current_date + timedelta(days=1 * 365)
    # Validity = new_date
    res=Workrequest.objects.filter(JOBCARD__USER__LOGIN_id=lid)
    # print(res)
    if res.exists():
        return JsonResponse({'status':'no'})
    else:
        obj=Workrequest()
        obj.Area=area
        obj.Taxreciept=path
        obj.Status='pending'
        obj.JOBCARD=Jobcard.objects.filter(USER__LOGIN_id=lid)[0]
        obj.Date=datetime.now().strftime('%Y-%m-%d')
        obj.SCHEME_id=sid
        obj.save()
        return JsonResponse({'status':'ok'})


def staffviewworkreqapply(request):
    res=Workrequest.objects.filter(Status='pending')
    return render(request,'staff/workrequestapplication.html',{'data':res})


def staffviewworkreqapply_post(request):
    fromdate=request.POST['textfield']
    todate=request.POST['textfield2']
    res=Workrequest.objects.filter(Status='pending',Date__range=[fromdate,todate])
    return render(request,'staff/workrequestapplication.html',{'data':res})

def approveworkreq(request,id):
    res=Workrequest.objects.filter(id=id).update(Status='approved')


    return HttpResponse(
        '''<script>alert('Verified');window.location='/myapp/staffviewworkreqapply/'</script>''')


def rejectworkreq(request,id):
    res=Workrequest.objects.filter(id=id).update(Status='rejected')
    return HttpResponse(
        '''<script>alert('Rejected');window.location='/myapp/staffviewworkreqapply/'</script>''')


def staff_viewdistinctwrkrequest(request):
    res = Forwardedwrkrequest.objects.filter(WORKREQUEST__Status="approved")
    l = []
    l2=[]
    for i in res:
        if i.EMPLOYEEMATE.id in l2:
            continue
        l2.append(i.EMPLOYEEMATE.id)
        if Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.EMPLOYEEMATE.EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD_id):
            mm = Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.EMPLOYEEMATE.EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD_id)[0]
            l.append({'id': i.id, 'workid':i.WORKREQUEST_id, 'emate':i.EMPLOYEEMATE.EMPLOYEEMATEREQ.JOBCARDMEMBERS.Name,'ward':i.EMPLOYEEMATE.EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD})
    return render(request, 'staff/viewdistinctworkrequest.html', {'data': l})


# def staff_viewworkrequest(request):
def staff_viewworkrequest(request, id):
    res=Workrequest.objects.filter(Status='approved', JOBCARD__USER__WARD_id=id)
    l=[]
    for i in res:
        if Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.JOBCARD.USER.WARD_id):
            mm = Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.JOBCARD.USER.WARD_id)[0]
            l.append({'id':i.id,'wid':i.JOBCARD.USER.WARD.id,'Wnum':i.JOBCARD.USER.WARD.Wnum,'Wname':i.JOBCARD.USER.WARD.Wname,'uu':mm.EMPLOYEEMATEREQ.USER.Uname})
    return render(request,'staff/viewempworkrequest.html',{'data':l})

def staff_viewworkrequestverified(request):
    res=Workrequest.objects.filter(Status='verified')
    l=[]
    for i in res:
        if Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.JOBCARD.USER.WARD_id):
            mm = Employeemate.objects.filter(EMPLOYEEMATEREQ__USER__WARD_id=i.JOBCARD.USER.WARD_id)[0]
            l.append({'id':i.id,'wid':i.JOBCARD.USER.WARD.id,'Wnum':i.JOBCARD.USER.WARD.Wnum,'Wname':i.JOBCARD.USER.WARD.Wname,'uu':mm.EMPLOYEEMATEREQ.USER.Uname})
    return render(request,'staff/viewempworkrequestverify.html',{'data':l})


def approvedstaffviewworkreqapply(request,id):
    res=Workrequest.objects.filter(Status='approved', JOBCARD__USER__WARD_id=id)
    # res=Workrequest.objects.filter(Status='approved')
    wards = Workrequest.objects.filter(Status='approved')[0].JOBCARD.USER.WARD.id
    l = []
    resEmcount=Employee.objects.filter(JOBCARDMEMBERS__JOBCARD__USER__WARD_id=id).count()
    emps = 0
    days = 0
    sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0


    for i in res:
        #  100 sqm = 10 workers
        total_area = float(i.Area)
        # print(total_area)
        area_per_worker = 100  # 100 sqm per 10 workers
        # print(total_workers)
        area_covered_per_day_per_worker = 10  # each worker covers 10 sqm per day

        total_workers = round(total_area / area_per_worker)

        total_area_covered_per_day = total_workers * area_covered_per_day_per_worker
        total_days = round(total_area / total_area_covered_per_day)

        emps += total_workers
        days += total_days
        print(total_workers, total_days, i.Area)

        l.append({
            'Date': i.Date,
            'Area': i.Area,
            'workers':total_workers ,
            'Cardno': i.JOBCARD.Cardno,
            'Head': i.JOBCARD.Head,
            'Days': total_days,

        })
    print(emps, resEmcount)
    return render(request,'staff/workrequestapplication.html',{
        'data':l,
        'wrkreq': res,
        'emp':int(emps),
        # 'emp':str(emps),
        'area':sum_of_areas,
        'wid':id,
        'resEmcount':int(resEmcount),
        # 'resEmcount':str(resEmcount),
        'days':days
    })




#
#
# def approvedstaffviewworkreqapply(request,id):
#     res=Workrequest.objects.filter(Status='approved',id=id)
#     wards = Workrequest.objects.filter(Status='approved', id=id)[0].JOBCARD.USER.WARD.id
#     l = []
#     resEmcount=Employee.objects.filter(JOBCARDMEMBERS__JOBCARD__USER__WARD_id=wards).count()
#     emps = Jobcardmembers.objects.exclude()
#     sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0
#
#
#     for i in res:
#
#         work_request = i
#         total_work_area_sqm = work_request.Area
#
#         num_workers = Employee.objects.exclude(JOBCARDMEMBERS__JOBCARD_id=work_request.JOBCARD_id).count()
#
#         if num_workers > 0 and total_work_area_sqm > 0:
#             worker_density = num_workers / float(total_work_area_sqm)
#             workers_needed = total_work_area_sqm * 0.5
#             print(worker_density, num_workers, total_work_area_sqm, workers_needed)
#
#             l.append({
#                 'Date': i.Date,
#                 'Area': i.Area,
#                 'Cardno': i.JOBCARD.Cardno,
#                 'Head': i.JOBCARD.Head,
#
#             })
#         print(l,'llll')
#     return render(request,'staff/workrequestapplication.html',{'data':l,'emp':str(len(emps)),'area':sum_of_areas,'wid':id, 'resEmcount':str(resEmcount)})
#     # return render(request,'staff/workrequestapplication.html',{'data':l,'emp':str(len(emps)),'area':sum_of_areas,'wid':id, 'resEmcount':str(resEmcount)})
#






#
# def workallocation(request,id):
#     # 10cent area=5workers
#     obj=Workrequest.objects.get(id=id)
#     total_work_area=obj.Area #sqm
#
#
#     res=Employee.objects.filter()
#     num_workers=res.count()
#
#     return render(request,'staff/viewworkers.html',{'data':res, 'id':id})

#
# def workallocation(request,id):
#     work_request = Workrequest.objects.get(id=id)
#     total_work_area_sqm = work_request.Area
#     res = Employee.objects.filter()
#     num_workers = res.count()
#     if num_workers > 0 and total_work_area_sqm > 0:
#         worker_density = num_workers / total_work_area_sqm
#         return worker_density
#     else:
#         return render(request,'staff/viewworkers.html',{'data':res, 'id':id})
#






def workallocation(request, id):
    work_request = Workrequest.objects.get(id=id)
    total_work_area_sqm = work_request.Area
    emps = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id)
    num_workers = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id).count()
    if num_workers > 0 and total_work_area_sqm > 0:
        worker_density = num_workers / float(total_work_area_sqm)
        workers_needed = total_work_area_sqm * 0.5
        print(worker_density, num_workers, total_work_area_sqm, workers_needed)
        return render(request, 'staff/viewworkers.html',{
                       'data': emps[:workers_needed],
                       'id': id,
                       'worker_density': str(worker_density),
                       })
    else:
        return render(request, 'staff/viewworkers.html', {'data': emps, 'id': id})





def workallocationpost(request):
    wid=request.POST['wid']
    emp=request.POST.getlist('selected_items')


    for i in emp:

        res=WorkAllocation()
        res.Status='assigned'
        from datetime import datetime
        res.Date=datetime.now().strftime('%Y-%m-%d')
        res.WORKREQUEST_id=wid
        res.JOBCARDMEMMBERS_id=i
        res.save()
    return HttpResponse(
        '''<script>alert('assigned');window.location='/myapp/staffHome/'</script>''')





def approvedstaffviewworkreqapply_post(request):
    fromdate = request.POST['textfield']
    todate = request.POST['textfield2']
    res=Workrequest.objects.filter(Status='approved',Date__range=[fromdate,todate])
    return render(request,'staff/distinctworkrequestapplication.html',{'data':res})


def rejecteddstaffviewworkreqapply(request):
    res=Workrequest.objects.filter(Status='rejected')
    return render(request,'staff/rejectedworkrequestapplication.html',{'data':res})

def rejecteddstaffviewworkreqapply_post(request):
    fromdate = request.POST['textfield']
    todate = request.POST['textfield2']
    res=Workrequest.objects.filter(Status='rejected',Date__range=[fromdate,todate])
    return render(request,'staff/rejectedworkrequestapplication.html',{'data':res})


def employeemateviewpublishlist(request):
    res=Workrequest.objects.filter(Status='verified')
    l=[]
    for i in res:
        if WorkAllocation.objects.filter(WORKREQUEST_id=i.id).exists():
            l.append({
            'id':i.id,
            'Date':i.Date,
            'Status':i.Status,
            'Area':i.Area,
            'name':i.JOBCARD.Head,
            'hname':i.JOBCARD.Cardno,
            'Sdate':i.Sdate,
            'Edate':i.Edate,
            'wardno':i.JOBCARD.USER.WARD.Wnum,

                  })
    return JsonResponse({'status':'ok','data':l})



def markattendance(request):
    WORKREQUEST=request.POST['sid']
    print(WORKREQUEST,'sssssss')
    res=WorkAllocation.objects.filter(WORKREQUEST_id=WORKREQUEST,Status='allocated')
    # res=Employee.objects.filter()
    l = []
    ex = []

    for i in res:
        if i.JOBCARDMEMMBERS.id in ex:
            continue
        if Attendance.objects.filter(WORKREQUEST_id=i.WORKREQUEST_id,Date=datetime.now().date().today(), JOBCARDMEMMBERS=i.JOBCARDMEMMBERS).exists():
            continue

        ex.append(i.JOBCARDMEMMBERS.id)
        l.append({
            'id': i.id,

            'name': i.JOBCARDMEMMBERS.Name,
            'head': i.JOBCARDMEMMBERS.JOBCARD.Head,

        })
    return JsonResponse({'status': 'ok', 'data': l})

def emp_add_attendance(request):
    lid=request.POST['lid']
    wid=request.POST['wid']
    Employee=request.POST['name']
    sk=Employee.split(',')
    for i in sk:
        print(i)
        ms=Attendance()
        ms.EMPLOYEEMATE=Employeemate.objects.get(LOGIN_id=lid)
        ms.JOBCARDMEMMBERS_id=WorkAllocation.objects.get(id=i).JOBCARDMEMMBERS_id
        ms.Status='Present'
        ms.Date=datetime.now().strftime('%Y-%m-%d')
        ms.WORKREQUEST_id=wid
        ms.save()

    else:
        ww=WorkAllocation.objects.filter(WORKREQUEST_id=wid)
        for i in ww:
            if Attendance.objects.filter(JOBCARDMEMMBERS_id = i.JOBCARDMEMMBERS_id, WORKREQUEST_id=i.WORKREQUEST_id, Date=datetime.now().date()).exists():
                continue
            ms = Attendance()
            ms.EMPLOYEEMATE = Employeemate.objects.get(LOGIN_id=lid)
            ms.JOBCARDMEMMBERS_id = WorkAllocation.objects.get(id=i.id).JOBCARDMEMMBERS_id
            ms.Status = 'Absent'
            ms.Date = datetime.now().strftime('%Y-%m-%d')
            ms.WORKREQUEST_id=wid
            ms.save()
    return JsonResponse({'status': "ok"})





def employeemateviewworkrequest(request):
    lid=request.POST['lid']
    print(lid)
    res=Workrequest.objects.filter(Status='pending',JOBCARD__USER__WARD_id=Employeemate.objects.get(LOGIN_id=lid).EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD_id)
    l=[]
    for i in res:
        l.append({
            'id':i.id,
            'Date':i.Date,
            'Area':i.Area,
            'head':i.JOBCARD.Head,
            'cardno':i.JOBCARD.Cardno,
                  })
    return JsonResponse({'status':'ok','data':l})
#######################
def employeematedeletewrkrequest(request):
    id=request.POST['id']
    Workrequest.objects.filter(id=id).delete()
    return JsonResponse({'status':"ok"})


def employeeallocatework(request):
    employ=request.POST['emid']
    wrkrqst=request.POST['wid']
    cc=WorkAllocation()
    from datetime import datetime
    cc.Date=datetime.now()
    cc.WORKREQUEST_id=wrkrqst
    cc.JOBCARDMEMMBERS=employ
    cc.Status='Pending'
    cc.save()
    return JsonResponse({'status':'ok'})

def employeemateviewallocatedarea(request):
    res=WorkAllocation.objects.all()
    l=[]
    for i in res:
        l.append({
            'id':i.id,
            'Date':i.Date,
            'Area':i.WORKREQUEST.Area,
            'headname':i.JOBCARDMEMMBERS.JOBCARD.Head,
            'Status':i.Status,
                  })
    return JsonResponse({'status':'ok','data':l})

def employeemateviewworkhistory(request):
    res=WorkAllocation.objects.filter()
    l=[]
    for i in res:
        l.append({
            'id':i.id,
            'Date':i.Date,
            'Area':i.WORKREQUEST.Area,
            'headname':i.JOBCARDMEMMBERS.JOBCARD.Head,
            'Status':i.Status,
                  })
    return JsonResponse({'status':'ok','data':l})

def employeematefoewardtostaff(request):
    rid=request.POST['rid']
    lid=request.POST['lid']
    obj=Workrequest.objects.filter(id=rid).update(Status='approved')
    obj2=Forwardedwrkrequest()
    obj2.WORKREQUEST_id=rid
    obj2.EMPLOYEEMATE=Employeemate.objects.get(LOGIN_id=lid)
    obj2.save()
    return JsonResponse({'status':'ok'})



#
# def employeemateaddworkallocation(request):
#
#
#     # 1 person = 10 squaremeter
#     # total work area
#     # total workers
#
#
#
#     #
#     res=total_work_area_sqm/(total workers*100(squaremeter))
#
#
#     id=request.POST['id']
#     work_request = Workrequest.objects.get(id=id)
#     total_work_area_sqm = work_request.Area
#     emps = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id)
#     num_workers = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id).count()
#     if num_workers > 0 and total_work_area_sqm > 0:
#         worker_density = num_workers / float(total_work_area_sqm)
#         workers_needed = total_work_area_sqm * 0.5
#         print(worker_density, num_workers, total_work_area_sqm, workers_needed)
#         # return render(request, 'staff/viewworkers.html',{
#         #                'data': emps[:workers_needed],
#         #                'id': id,
#         #                'worker_density': str(worker_density),
#         #                })
#
#
#
#
#         return JsonResponse({'status': 'ok','data':l})
#
#     else:
#         # return render(request, 'staff/viewworkers.html', {'data': emps, 'id': id})
#         return JsonResponse({'status': 'no'})
#




        # 1 person = 10 squaremeter
        # total work area
        # total workers

        # d=a/(w*10)


def employeemateaddworkallocation(request):
    lid=request.POST['lid']
    res = Workrequest.objects.filter(Status='verified')
    l = []
    # emps = Jobcardmembers.objects.exclude()
    # sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0
    #
    # for work_request in res:
    #     total_work_area_sqm = work_request.Area
    #     num_workers = Employee.objects.exclude(JOBCARDMEMBERS__JOBCARD_id=work_request.JOBCARD_id).count()
    #
    #     if num_workers > 0 and total_work_area_sqm > 0:
    #         worker_density = (total_work_area_sqm / num_workers) * 100
    #         days_needed = total_work_area_sqm / num_workers
    #
    #         l.append({
    #             'id': work_request.id,
    #             'Date': work_request.Date,
    #             'Area': work_request.Area,
    #             'Cardno': work_request.JOBCARD.Cardno,
    #             'Head': work_request.JOBCARD.Head,
    #             'WorkerDensity': worker_density,
    #             'DaysNeeded': str(days_needed)[:1]
    #         })
    # response_data = {
    #     'status': 'ok',
    #     'data': l,
    #     'emp': len(emps),
    #     'area': sum_of_areas
    # }

    res = Workrequest.objects.filter(Status='verified', JOBCARD__USER__WARD_id=Employeemate.objects.get(LOGIN_id=lid).EMPLOYEEMATEREQ.JOBCARDMEMBERS.JOBCARD.USER.WARD.id)
    # res = Workrequest.objects.filter(Status='verified', JOBCARD__USER__WARD_id=id)
    # res=Workrequest.objects.filter(Status='approved')
    wards = Workrequest.objects.filter(Status='verified')[0].JOBCARD.USER.WARD.id
    l = []
    resEmcount = Employee.objects.filter(JOBCARDMEMBERS__JOBCARD__USER__WARD_id=wards).count()
    emps = 0
    days = 0
    sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0

    for i in res:
        if WorkAllocation.objects.filter(WORKREQUEST_id=i.id).exists():
            continue
        #  100 sqm = 10 workers
        total_area = float(i.Area)
        # print(total_area)
        area_per_worker = 100  # 100 sqm per 10 workers
        # print(total_workers)
        area_covered_per_day_per_worker = 10  # each worker covers 10 sqm per day

        total_workers = round(total_area / area_per_worker)

        total_area_covered_per_day = total_workers * area_covered_per_day_per_worker
        total_days = round(total_area / total_area_covered_per_day)

        emps += total_workers
        days += total_days
        print(total_workers, total_days, i.Area)

        l.append({
            'id': i.id,
            'Date': i.Date,
            'Area': i.Area,
            'WorkerDensity': total_workers,
            'Cardno': i.JOBCARD.Cardno,
            'Head': i.JOBCARD.Head,
            'DaysNeeded': total_days,
            'Sdate': i.Sdate,
            'Edate': i.Edate,})

    response_data = {
        'status': 'ok',
        'data': l,
        'emp': emps,
        'area': sum_of_areas
    }
    return JsonResponse(response_data)
















#
#
# def employeemateallocateworkerstoworker(request):
#     res=Workrequest.objects.filter(Status='approved')
#     l = []
#     emps = Jobcardmembers.objects.exclude()
#     sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0
#
#
#     for i in res:
#
#         work_request = i
#         total_work_area_sqm = work_request.Area
#
#         num_workers = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id).count()
#
#         if num_workers > 0 and total_work_area_sqm > 0:
#             worker_density = num_workers / float(total_work_area_sqm)
#             workers_needed = total_work_area_sqm * 0.5
#             print(worker_density, num_workers, total_work_area_sqm, workers_needed)
#
#             l.append({
#                 'id': i.id,
#                 'Area': i.Area,
#                 'Cardno': i.JOBCARD.Cardno,
#                 'Head': i.JOBCARD.Head,
#
#             })
#         print(l,'llll')
#     return JsonResponse({'status':'ok','data':l})
#







def employeemateallocateworkerstoworker(request):
    res=WorkAllocation.objects.filter(WORKREQUEST__Status='approved')
    l = []
    emps = Jobcardmembers.objects.exclude()
    sum_of_areas = res.aggregate(total_area=Sum('Area'))['total_area'] or 0


    for i in res:

        work_request = i
        total_work_area_sqm = work_request.Area

        num_workers = Jobcardmembers.objects.exclude(JOBCARD_id=work_request.JOBCARD_id).count()

        if num_workers > 0 and total_work_area_sqm > 0:
            worker_density = num_workers / float(total_work_area_sqm)
            workers_needed = total_work_area_sqm * 0.5
            print(worker_density, num_workers, total_work_area_sqm, workers_needed)

            l.append({
                'id': i.id,
                'Area': i.Area,
                'Cardno': i.JOBCARD.Cardno,
                'Head': i.JOBCARD.Head,

            })
        print(l,'llll')
    return JsonResponse({'status':'ok','data':l})



def viewassignedworkers(request):
    id=request.POST['rid']

    work_request = Workrequest.objects.get(id=id)
    i = work_request
    total_work_area_sqm = work_request.Area
    # emps = Employee.objects.exclude(JOBCARD_id=work_request.JOBCARD_id)
    num_workers = Employee.objects.exclude(JOBCARDMEMBERS__JOBCARD_id=work_request.JOBCARD_id).count()
    workers_needed=0
    if num_workers > 0 and total_work_area_sqm > 0:
        worker_density = num_workers / float(total_work_area_sqm)
        workers_needed = total_work_area_sqm * 0.5

    # res=Employee.objects.filter(Status='verified')[:workers_needed]
    res=Employee.objects.filter(Status='verified', JOBCARDMEMBERS__JOBCARD__USER__WARD_id=i.JOBCARD.USER.WARD.id)
    l = []
    ex = []
    for i in res:
        if i.JOBCARDMEMBERS.id in ex:
            continue
        # if WorkAllocation.objects.filter(WORKREQUEST_id=id, JOBCARDMEMMBERS=i.JOBCARDMEMBERS).exists():
        if WorkAllocation.objects.filter(JOBCARDMEMMBERS=i.JOBCARDMEMBERS).exists():
            continue
        ex.append(i.JOBCARDMEMBERS.id)
        l.append({
            'id': i.JOBCARDMEMBERS.id,
            'name': i.JOBCARDMEMBERS.Name,
            # 'head': i.JOBCARDMEMMBERS.JOBCARD.Head,

        })
    print(l,'lllllllllllllllllllll')
    return JsonResponse({'status': 'ok', 'data': l})


def emp_add_allocation(request):
    wid=request.POST['wid']
    Jobcardmembers=request.POST['name']
    sk=Jobcardmembers.split(',')

    # if WorkAllocation.objects.filter(JOBCARDMEMMBERS_id=sk,Status='allocated').exists():
    #     return JsonResponse({'status':'no'})

    for i in sk:
        print(i)
        ms=WorkAllocation()
        ms.WORKREQUEST_id=wid
        ms.JOBCARDMEMMBERS_id=i
        ms.Status='allocated'
        ms.Date=datetime.now().strftime('%Y-%m-%d')
        ms.save()
    return JsonResponse({'status': "ok"})



def usersendworkrequest(request):
    lid=request.POST['lid']
    mid=request.POST['mid']
    jjj=Jobcardmembers.objects.filter(JOBCARD__USER__LOGIN_id=lid)[0]
    # if Employee.objects.filter(JOBCARDMEMBERS_id=mid, Date=datetime.now().strftime('%Y-%m-%d')).exists():
    if Employee.objects.filter(JOBCARDMEMBERS_id=mid).exists():
        return JsonResponse({'status': "no"})

    res=Employee()

    res.JOBCARDMEMBERS_id=mid
    res.Date=datetime.now().strftime('%Y-%m-%d')
    res.Status='pending'
    res.save()

    return JsonResponse({'status': "ok"})



def staff_viewemployeesreq(request,id):
    # res=Employee.objects.filter(JOBCARDMEMBERS__)
    res=Employee.objects.filter(JOBCARDMEMBERS__JOBCARD__USER__WARD_id=id)
    l=[]
    for i in res:

        l.append({'id':i.id,'Name':i.JOBCARDMEMBERS.Name,"Head":i.JOBCARDMEMBERS.JOBCARD.Head,'Cardno':i.JOBCARDMEMBERS.JOBCARD.Cardno})
    return render(request,'staff/viewworkerslist.html',{'data':l})




def staff_approvewrkrequest(request,id):
    obj=Workrequest.objects.filter(JOBCARD__USER__WARD_id=id,Status='approved')
    l = []
    emps = 0
    days = 0

    for i in obj:
        #  100 sqm = 10 workers
        total_area = float(i.Area)
        # print(total_area)
        area_per_worker = 100  # 100 sqm per 10 workers
        # print(total_workers)
        area_covered_per_day_per_worker = 10  # each worker covers 10 sqm per day

        total_workers = round(total_area / area_per_worker)

        total_area_covered_per_day = total_workers * area_covered_per_day_per_worker
        total_days = round(total_area / total_area_covered_per_day)

        emps += total_workers
        days += total_days
        # print(total_workers, total_days, i.Area)

        from django.utils import timezone
        from datetime import timedelta

        current_date = timezone.now()

        # Add 2 days to the current date
        start_date = current_date + timedelta(days=2)

        # Convert future_date to the desired format (e.g., string)
        start_date_str = start_date.strftime("%Y-%m-%d")

        # Add 10 days to the future date
        end_date = start_date + timedelta(days=total_days)

        # Convert future_date_plus_10_days to the desired format (e.g., string)
        end_date_str = end_date.strftime("%Y-%m-%d")
        print(end_date_str,start_date_str)
        i.Sdate=start_date_str
        i.Edate=end_date_str
        i.save()


    res=Workrequest.objects.filter(JOBCARD__USER__WARD_id=id).update(Status='verified')
    # res=Workrequest.objects.filter(id=id).update(Status='verified')
    res2=Employee.objects.filter(JOBCARDMEMBERS__JOBCARD__USER__WARD_id=id).update(Status='verified')
    return HttpResponse(
        '''<script>alert('Verified');window.location='/myapp/staff_viewdistinctwrkrequest/'</script>''')





def staff_viewattendance(request,id):
# def staff_viewattendance(request):
    res=Attendance.objects.filter(WORKREQUEST_id=id)
    # res=Attendance.objects.all()
    return render(request,'staff/view_attendance.html',{'data':res})


def userviewworkhistory(request):
    lid = request.POST['lid']
    res=WorkAllocation.objects.filter(JOBCARDMEMMBERS__JOBCARD__USER__LOGIN_id=lid)
    l=[]
    for i in res:
        l.append({
            'id':i.id,
            'Date':i.Date,
            'Area':i.WORKREQUEST.Area,
            'headname':i.JOBCARDMEMMBERS.Name,
            'Status':i.Status,
                  })
    return JsonResponse({'status':'ok','data':l})


def user_view_allocation(request):
    lid=request.POST['lid']
    # mid=request.POST['mid']
    l=[]
    a=WorkAllocation.objects.filter(JOBCARDMEMMBERS__JOBCARD__USER__LOGIN_id=lid)
    for i in a:
        l.append({'id':i.JOBCARDMEMMBERS.id,'Date':i.Date,
                  'Status':i.Status,
                  'wid':i.WORKREQUEST.id,
                  'Area':i.WORKREQUEST.Area,
                  'Head':i.WORKREQUEST.JOBCARD.Head,
                  'CardNo':i.WORKREQUEST.JOBCARD.Cardno,
                  'employee': i.JOBCARDMEMMBERS.Name,
                  'Sdate':i.WORKREQUEST.Sdate,
                  'Edate':i.WORKREQUEST.Edate,'jmid':i.JOBCARDMEMMBERS.id,
                  })
    return JsonResponse({'status': 'ok', 'data': l})

#
# def view_my_attendance(request):
#     lid = request.POST['lid']
#     wid = request.POST['wid']
#     l=[]
#     print(lid)
#     a = Attendance.objects.filter(JOBCARDMEMMBERS_id=lid, WORKREQUEST_id=wid)
#     for i in a:
#         from datetime import datetime
#
#         date_string1 = i.WORKREQUEST.Sdate
#         date_string2 = i.WORKREQUEST.Edate
#
#         date1 = datetime.strptime(date_string1, "%Y-%m-%d")
#         date2 = datetime.strptime(date_string2, "%Y-%m-%d")
#
#         if date1 > date2:
#             date1, date2 = date2, date1
#
#         date_difference = date2 - date1
#
#         days_difference = date_difference.days
#
#         ds = []
#         for ij in range(days_difference + 1):
#             current_date = date1 + timedelta(days=ij)
#             status="pending"
#             if str(current_date.strftime("%Y-%m-%d"))==str(i.Date):
#                 status=i.Status
#             if current_date.strftime("%Y-%m-%d") in ds:
#                 continue
#             ds.append(current_date.strftime("%Y-%m-%d"))
#             l.append({'id':i.id,'Date':current_date.strftime("%Y-%m-%d"),
#                   'Status':status,'Days':ds,})
#     return JsonResponse({'status': 'ok', 'data': l})

def view_my_attendance(request):
    lid = request.POST['lid']
    wid = request.POST['wid']

    Sdate = request.POST['Sdate']
    Edate = request.POST['Edate']
    l = []
    # a = Attendance.objects.filter(JOBCARDMEMMBERS_id=lid, WORKREQUEST_id=wid)

    # Initialize ds outside the outer loop
    # ds = []

    # for attendance_record in a:
    #     from datetime import datetime, timedelta
    #
    #     date_string1 = attendance_record.WORKREQUEST.Sdate
    #     date_string2 = attendance_record.WORKREQUEST.Edate
    #
    #     date1 = datetime.strptime(date_string1, "%Y-%m-%d")
    #     date2 = datetime.strptime(date_string2, "%Y-%m-%d")
    #
    #     if date1 > date2:
    #         date1, date2 = date2, date1
    #
    #     date_difference = date2 - date1
    #     days_difference = date_difference.days
    #
    #     for day_offset in range(days_difference + 1):
    #         current_date = date1 + timedelta(days=day_offset)
    #         status = "pending"
    #
    #         if str(current_date.strftime("%Y-%m-%d")) == str(attendance_record.Date):
    #             status = attendance_record.Status
    #
    #         if current_date.strftime("%Y-%m-%d") in ds:
    #             continue
    #
    #         ds.append(current_date.strftime("%Y-%m-%d"))
            # l.append({
            #     'id': attendance_record.id,
            #     'Date': current_date.strftime("%Y-%m-%d"),
            #     'Status': status,
            # })

    a = Attendance.objects.filter(JOBCARDMEMMBERS_id=lid,Date__range=[Sdate,Edate], WORKREQUEST_id=wid)
    for i in a:
        l.append({"id":i.id,"Status":i.Status,
                                'Date':i.Date})
    print(a,l)
    return JsonResponse({'status': 'ok', 'data': l})

# from datetime import datetime, timedelta
#
# def view_my_attendance(request):
#     lid = request.POST['lid']
#     Sdate = request.POST['Sdate']
#     Edate = request.POST['Edate']
#
#     print(Sdate)
#     print(Edate)
#     attendance_dict = [] # Use a dictionary to store unique dates and their statuses
#     a = Attendance.objects.filter(JOBCARDMEMMBERS__JOBCARD__USER__LOGIN_id=lid,Date__range=[Sdate,Edate])
#     # for i in a:
#     #     attendance_dict.append({"id":i.id,"Status":i.Status,
#     #                             'Date':i.Date})
#
#     for i in a:
#         date_string1 = i.WORKREQUEST.Sdate
#         date_string2 = i.WORKREQUEST.Edate
#
#         date1 = datetime.strptime(date_string1, "%Y-%m-%d")
#         date2 = datetime.strptime(date_string2, "%Y-%m-%d")
#
#         if date1 > date2:
#             date1, date2 = date2, date1
#
#         date_difference = date2 - date1
#         days_difference = date_difference.days
#
#         for ij in range(days_difference + 1):
#             current_date = date1 + timedelta(days=ij)
#
#             attendance_dict[current_date.strftime("%Y-%m-%d")] = 'pending'
#
#     for i in a:
#         status = i.Status
#         date = i.Date
#         attendance_dict[date.strftime("%Y-%m-%d")] = status  # Update status for dates where attendance is marked
#
#
#     return JsonResponse({'status': 'ok', 'data': attendance_dict })
